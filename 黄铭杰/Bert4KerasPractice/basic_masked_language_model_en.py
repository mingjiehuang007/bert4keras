#! -*- coding: utf-8 -*-
# 测试代码可用性: MLM
import random

import numpy as np
from bert4keras.models import build_transformer_model
from bert4keras.tokenizers import Tokenizer
from bert4keras.snippets import to_array
from docx import Document           # 三方库python-docx
from docx.shared import RGBColor


# 导入模型的设置，模型本体，字典
config_path = 'root\\kg\\bert\\uncased_L-12_H-768_A-12\\bert_config.json'
checkpoint_path = 'root\\kg\\bert\\uncased_L-12_H-768_A-12\\bert_model.ckpt'
dict_path = 'root\\kg\\bert\\uncased_L-12_H-768_A-12\\vocab.txt'
tokenizer = Tokenizer(dict_path, do_lower_case=True)  # 建立分词器


# 将test和生成的句子读入内存备用
def mask_sen(file_path_orig, file_path_unilm, output_filename):
    f_test = open(file_path_orig, "r", encoding="utf-8")
    f_pred = open(file_path_unilm, "r", encoding="utf-8")
    lines_test = f_test.readlines()
    lines_pred = f_pred.readlines()
    f_test.close()
    f_pred.close()

    sen_pair_text = []
    for i in range(len(lines_test)):
        sen1 = lines_test[i][0:lines_test[i].find("\t")]
        sen2 = lines_pred[i][0:-1]
        words1 = tokenizer.tokenize(sen1)
        words2 = tokenizer.tokenize(sen2)

        sen_masked = ""
        sen_pair = sen2 + "\t"
        for w in words2:
            if w != "[CLS]" and w != "[SEP]":
                if w in words1:
                    sen_masked += "[MASK] "
                else:
                    sen_masked += w + " "
        sen_pair += sen_masked.strip()
        sen_pair_text.append(sen_pair + "\n")

    with open(output_filename, "w", encoding="utf-8") as f_masked:
        for i in sen_pair_text:
            f_masked.write(i)
    f_masked.close()


def tokenize_file(origin_file_path, tokenized_file_path):
    f = open(origin_file_path, "r", encoding="utf-8")
    lines = f.readlines()
    f.close()

    result_list = []
    for line in lines:
        line_temp = line.split("\t")[0]
        arr = tokenizer.tokenize(line_temp)
        print(arr)
        arr_str = ""
        for i in arr:
            arr_str += i + " "

        result_list.append(line_temp + "\t" + arr_str.strip() )

    f_result = open(tokenized_file_path, "w", encoding="utf-8")
    for i in result_list:
        f_result.write(i + "\n")
    f_result.close()


# 对句子进行mask处理，并输出至文件
def create_masked_text():
    path_test = "datum/tg_post\\test.txt"
    path_50k = "datum/tg_post\\unilm_quora_50k_random_rep_0.5_0.3_8000_bleu_0.2731582345959125.txt"
    path_100k = "datum/tg_post\\unilm_quora_100k_random_rep_x_o_0.6_0.4_28000_bleu_0.28027037196011884.txt"
    path_150k = "datum/tg_post\\unilm_quora_150k_random_rep_x_o_0.5_0.3_524000_bleu_0.2903091303080638.txt"

    mask_sen(path_test, path_50k, "pretreatment\\50K_masked.txt")
    mask_sen(path_test, path_100k, "pretreatment\\100K_masked.txt")
    mask_sen(path_test, path_150k, "pretreatment\\150K_masked.txt")


# 获取列表中某元素的所有位置，保存至列表中
def get_indices(str_name, list_name):
    index_list = []
    n = 0
    for i in list_name:
        if i == str_name:
            index_list.append(n)
        n += 1
    return index_list


# 将mask好的文本送入BERT进行预测，并输出至文件
def create_mlm_file(masked_file_path, output_path):
    f = open(masked_file_path, "r", encoding="utf-8")
    lines = f.readlines()
    f.close()

    model = build_transformer_model(
        config_path=config_path, checkpoint_path=checkpoint_path, with_mlm=True
    )  # 建立模型，加载权重

    total_count = len(lines)
    count = 0
    result_list = []
    for line in lines:
        count += 1
        print(str(count) + "||" + str(total_count))

        sen_orig = line.strip().split("\t")[0]      # 获取原句
        sen_masked = line.strip().split("\t")[1]    # 获取mask后的句子
        words_orig = tokenizer.tokenize(sen_orig)
        words_masked = tokenizer.tokenize(sen_masked.replace("[MASK]","mmm"))   # 对mask后的句子进行分词

        token_ids, segment_ids = tokenizer.encode(sen_orig)                     # 对原句进行分词，对词编码
        batch_token_ids = token_ids
        batch_segment_ids = segment_ids

        position_list = get_indices("mmm", words_masked)                        # 获取被掩码词的位置
        for i in range(len(words_orig)):                                        # 对句子逐个单词进行循环
            if i not in position_list:                                          # 若当前单词未被掩码
                batch_token_ids[i] = token_ids[i]                               # 将该单词的id作为结果id
                batch_segment_ids[i] = segment_ids[i]
            else:                                                               # 若当前单词被掩码
                # n = random.randint(1,10)
                # p = 2
                # if n <= p+1:                                                    # 20%概率用原词
                #     batch_token_ids[i] = token_ids[i]                           # 将该单词的id作为结果id
                #     batch_segment_ids[i] = segment_ids[i]
                # else:                                                           # 80%概率掩码，并用BERT预测单词
                    token_ids_temp, segment_ids_temp = tokenizer.encode(sen_orig)   # 对原句进行分词，对词编码
                    token_ids_temp[i] = tokenizer._token_mask_id                    # 在BERT中对当前词进行掩码
                    token_ids_temp, segment_ids_temp = to_array([token_ids_temp], [segment_ids_temp])

                    probas = model.predict([token_ids_temp, segment_ids_temp])[0]   # 用mlm模型预测被mask掉的部分
                    mlm_temp = tokenizer.decode(probas[1:-1].argmax(axis=1))
                    token_ids_mlm, segment_ids_mlm = tokenizer.encode(mlm_temp)

                    if len(batch_token_ids) != len(token_ids_mlm):                  # 如果预测前后长度不一，则跳过
                        continue

                    batch_token_ids[i] = token_ids_mlm[i]                           # 将该单词的id作为结果id
                    batch_segment_ids[i] = segment_ids_mlm[i]

        batch_token_ids, batch_segment_ids = to_array([batch_token_ids], [batch_segment_ids])
        probas2 = model.predict([batch_token_ids,batch_segment_ids])[0]
        sen_output_str = tokenizer.decode(probas2[1:-1].argmax(axis=1))
        sen_output_str = sen_output_str.replace("?", " ?").replace(",", " ,").replace(".", " .").replace("'", " '")
        print(sen_orig)
        print(sen_output_str)
        result_list.append(sen_output_str)

    f_result = open(output_path, "w", encoding="utf-8")
    for i in result_list:
        f_result.write(i + "\n")
    f_result.close()


# 对句子逐个进行MASK，送入BERT进行预测，并输出至文件
def create_mlm_file_each(masked_file_path, output_path):
    f = open(masked_file_path, "r", encoding="utf-8")
    lines = f.readlines()
    f.close()

    model = build_transformer_model(
        config_path=config_path, checkpoint_path=checkpoint_path, with_mlm=True
    )  # 建立模型，加载权重

    total_count = len(lines)
    count = 0
    result_list = []

    for line in lines:
        count += 1
        print(str(count) + "||" + str(total_count))

        sen_orig = line.strip().split("\t")[0]      # 获取原句
        words_orig = tokenizer.tokenize(sen_orig)

        result_list.append("origin: " + sen_orig)
        sen_tokenized = ""
        for wordpiece in words_orig:
            sen_tokenized += wordpiece + " "
        result_list.append("tokenized: " + sen_tokenized.strip())

        for line in range(1,len(words_orig)-1):                                # 对句子逐个单词进行循环
            token_ids_temp, segment_ids_temp = tokenizer.encode(sen_orig)   # 对原句进行分词，对词编码
            token_ids_temp[line] = tokenizer._token_mask_id                    # 在BERT中对当前词进行掩码
            token_ids_temp, segment_ids_temp = to_array([token_ids_temp], [segment_ids_temp])

            probas = model.predict([token_ids_temp, segment_ids_temp])[0]   # 用mlm模型预测被mask掉的部分
            mlm_temp = tokenizer.decode(probas[1:-1].argmax(axis=1))
            mlm_temp = mlm_temp.replace("?", " ?").replace(",", " ,").replace(".", " .").replace("'", " '")
            print(line, end="|")
            print(len(words_orig)-2, end="\t")
            print(mlm_temp)
            result_list.append(mlm_temp)

    document = Document()                                   # 新建word文档
    position_in_tokens = 1
    for output_line in result_list:
        flag = False
        paragraph = document.add_paragraph()  # 新建段落
        color_blue = (0, 176, 240)  # 蓝色
        color_red = (255, 0, 0)  # 红色
        if output_line[0:6] == "origin" or output_line[0:9] == "tokenized":
            run = paragraph.add_run(output_line)  # 在段落中，新建语句
            run.font.color.rgb = RGBColor(*color_blue)  # 修改字体颜色
            position_in_tokens = 1
        else:
            output_line_tokens = tokenizer.tokenize(output_line)
            keyword = output_line_tokens[position_in_tokens]
            if "##" in keyword:
                keyword = keyword.replace("##", "")
                flag = True

            position_in_line = 0
            for i in range(1, position_in_tokens):
                if "##" in output_line_tokens[i]:
                    position_in_line += len(output_line_tokens[i]) - 2
                elif output_line_tokens[i] in ["'", ","]:
                    position_in_line += len(output_line_tokens[i])
                else:
                    position_in_line += len(output_line_tokens[i]) + 1

            if flag:
                line_left = output_line[0:position_in_line - 1]  # 用关键词切分当前文本
                line_mid = output_line[position_in_line - 1:position_in_line + len(keyword) - 1]
                line_right = output_line[position_in_line + len(keyword) - 1:]
            else:
                line_left = output_line[0:position_in_line]  # 用关键词切分当前文本
                line_mid = output_line[position_in_line:position_in_line + len(keyword)]
                line_right = output_line[position_in_line + len(keyword):]
            run = paragraph.add_run(line_left)
            color_temp = run.font.color.rgb
            run = paragraph.add_run(line_mid)
            run.font.color.rgb = RGBColor(*color_red)
            run = paragraph.add_run(line_right)
            run.font.color.rgb = color_temp

            position_in_tokens += 1

    document.save(output_path)                              # 保存文档

    # f_result = open(output_path, "w", encoding="utf-8")
    # for i in result_list:
    #     f_result.write(i + "\n")
    # f_result.close()


# create_masked_text()
# tokenize_file("pretreatment\\50K_masked.txt", "pretreatment\\50K_tokenized_file.txt")

# create_mlm_file("pretreatment\\50K_masked.txt", "result\\50K_result_80.txt")
# create_mlm_file("pretreatment\\100K_masked.txt", "result\\100K_result_80.txt")
# create_mlm_file("pretreatment\\150K_masked.txt", "result\\150K_result_80.txt")
create_mlm_file_each("pretreatment\\test_sentenses.txt", "result\\test_result.docx")
